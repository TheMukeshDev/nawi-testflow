"""
NAWI Sahayak — DOCX Report Generator

Generates editable DOCX test reports using python-docx.
Designed for technical/metrology reports that can be further edited.

Features:
- Structured sections with clear hierarchy
- Tables for observations and results
- OBSERVED / CALCULATED / LIMIT / RESULT color coding
- Headers and footers with page numbers
- Professional formatting
"""

from datetime import datetime, date
from io import BytesIO
from typing import Optional

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from .report_models import (
    TestReport, TestResult, TestProcedure, Observation,
    ReportIdentification, LaboratoryInfo, InstrumentInfo,
    TestConditions, TestEquipment, ComplianceResult,
    Signature, Attachment, ReportMetadata
)


# ============================================================================
# COLOR SCHEME (Professional)
# ============================================================================

class DocxColors:
    """Professional color palette."""
    
    # Header backgrounds
    HEADER_BG = RGBColor(0x1a, 0x36, 0x5d)    # Dark blue
    SECTION_BG = RGBColor(0x2c, 0x52, 0x82)   # Medium blue
    TABLE_HEADER = RGBColor(0x2d, 0x37, 0x48)  # Dark gray
    ALT_ROW = RGBColor(0xf7, 0xfa, 0xfc)       # Very light gray
    
    # Data type indicators
    OBSERVED_BG = RGBColor(0xeb, 0xf8, 0xff)   # Light blue
    CALCULATED_BG = RGBColor(0xf0, 0xff, 0xf4) # Light green
    LIMIT_BG = RGBColor(0xff, 0xf5, 0xf5)      # Light red
    RESULT_BG = RGBColor(0xfa, 0xf5, 0xff)     # Light purple
    
    # Status
    PASS_BG = RGBColor(0xc6, 0xf6, 0xd5)       # Light green
    FAIL_BG = RGBColor(0xfe, 0xd7, 0xd7)       # Light red
    
    # Text
    TEXT_PRIMARY = RGBColor(0x1a, 0x20, 0x2c)
    TEXT_SECONDARY = RGBColor(0x4a, 0x55, 0x68)
    WHITE = RGBColor(0xff, 0xff, 0xff)


# ============================================================================
# DOCX GENERATOR
# ============================================================================

class ReportDOCXGenerator:
    """
    Generates editable DOCX test reports using python-docx.
    
    Output can be further edited in Microsoft Word or similar applications.
    """
    
    def __init__(self):
        self.doc = None
    
    def generate(self, report: TestReport) -> bytes:
        """
        Generate DOCX report.
        
        Args:
            report: Complete test report data
            
        Returns:
            DOCX file as bytes
        """
        self.doc = Document()
        
        # Set default font
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(10)
        
        # Build report
        self._build_header(report)
        self._build_laboratory_info(report.laboratory)
        self._build_instrument_info(report.instrument)
        self._build_conditions(report.conditions)
        self._build_equipment(report.equipment)
        self._build_test_results(report.results)
        
        if report.compliance:
            self._build_compliance(report.compliance)
        
        if report.remarks:
            self._build_remarks(report.remarks)
        
        self._build_signatures(report)
        self._build_disclaimer()
        
        # Add header and footer
        self._add_header_footer(report)
        
        # Save to buffer
        buffer = BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        
        return buffer.getvalue()
    
    def _build_header(self, report: TestReport):
        """Build report header."""
        # Title
        title = self.doc.add_heading('TEST REPORT', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Report identification table
        table = self.doc.add_table(rows=3, cols=4)
        table.style = 'Table Grid'
        
        # Row 1
        self._set_cell(table, 0, 0, 'Report Number:', bold=True)
        self._set_cell(table, 0, 1, report.identification.report_number)
        self._set_cell(table, 0, 2, 'Date:', bold=True)
        self._set_cell(table, 0, 3, report.identification.report_date.strftime("%d %B %Y"))
        
        # Row 2
        self._set_cell(table, 1, 0, 'Standard:', bold=True)
        self._set_cell(table, 1, 1, report.identification.standard)
        self._set_cell(table, 1, 2, 'Version:', bold=True)
        self._set_cell(table, 1, 3, report.identification.standard_version)
        
        # Row 3
        self._set_cell(table, 2, 0, 'Revision:', bold=True)
        self._set_cell(table, 2, 1, report.identification.revision)
        
        self.doc.add_paragraph()
    
    def _build_laboratory_info(self, lab: LaboratoryInfo):
        """Build laboratory information section."""
        self.doc.add_heading('1. LABORATORY INFORMATION', level=1)
        
        table = self.doc.add_table(rows=7, cols=2)
        table.style = 'Table Grid'
        
        data = [
            ('Laboratory Name', lab.name),
            ('Address', '{} {} {} {}'.format(lab.address, lab.city, lab.state, lab.country)),
            ('Phone', lab.phone),
            ('Email', lab.email),
            ('Accreditation Body', lab.accreditation_body),
            ('Accreditation Number', lab.accreditation_number),
            ('Accreditation Expiry', lab.accreditation_expiry.strftime("%d %B %Y") if lab.accreditation_expiry else 'N/A'),
        ]
        
        for i, (label, value) in enumerate(data):
            self._set_cell(table, i, 0, label, bold=True)
            self._set_cell(table, i, 1, value)
        
        self.doc.add_paragraph()
    
    def _build_instrument_info(self, instrument: InstrumentInfo):
        """Build instrument information section."""
        self.doc.add_heading('2. INSTRUMENT INFORMATION', level=1)
        
        # Manufacturer
        self.doc.add_heading('2.1 Manufacturer', level=2)
        
        table = self.doc.add_table(rows=3, cols=2)
        table.style = 'Table Grid'
        
        mfr_data = [
            ('Manufacturer', instrument.manufacturer.name),
            ('Country', instrument.manufacturer.country),
            ('Address', instrument.manufacturer.address),
        ]
        
        for i, (label, value) in enumerate(mfr_data):
            self._set_cell(table, i, 0, label, bold=True)
            self._set_cell(table, i, 1, value)
        
        self.doc.add_paragraph()
        
        # Instrument
        self.doc.add_heading('2.2 Instrument', level=2)
        
        table = self.doc.add_table(rows=9, cols=2)
        table.style = 'Table Grid'
        
        inst_data = [
            ('Model Name', instrument.model_name),
            ('Model Number', instrument.model_number),
            ('Serial Number', instrument.serial_number),
            ('Instrument Type', instrument.instrument_type),
            ('Instrument Class', instrument.instrument_class),
            ('Maximum Capacity', '{} {}'.format(instrument.max_capacity, instrument.max_capacity_unit)),
            ('Minimum Capacity', '{} {}'.format(instrument.min_capacity, instrument.min_capacity_unit)),
            ('Scale Interval (d)', '{} {}'.format(instrument.scale_interval, instrument.scale_interval_unit)),
        ]
        
        if instrument.verification_scale_interval:
            inst_data.append(('Verification Scale Interval (e)', 
                            '{} {}'.format(instrument.verification_scale_interval, instrument.verification_scale_interval_unit)))
        
        for i, (label, value) in enumerate(inst_data):
            self._set_cell(table, i, 0, label, bold=True)
            self._set_cell(table, i, 1, value)
        
        self.doc.add_paragraph()
    
    def _build_conditions(self, conditions: TestConditions):
        """Build test conditions section."""
        self.doc.add_heading('3. TEST CONDITIONS', level=1)
        
        # Environmental conditions
        if conditions.conditions:
            table = self.doc.add_table(rows=len(conditions.conditions) + 1, cols=5)
            table.style = 'Table Grid'
            
            # Header
            headers = ['Parameter', 'Value', 'Unit', 'Range', 'Status']
            for i, header in enumerate(headers):
                self._set_cell(table, 0, i, header, bold=True, bg_color=DocxColors.TABLE_HEADER, font_color=DocxColors.WHITE)
            
            # Data
            for i, cond in enumerate(conditions.conditions, 1):
                range_str = ''
                if cond.min_value is not None and cond.max_value is not None:
                    range_str = '{} - {}'.format(cond.min_value, cond.max_value)
                
                self._set_cell(table, i, 0, cond.parameter)
                self._set_cell(table, i, 1, str(cond.value))
                self._set_cell(table, i, 2, cond.unit)
                self._set_cell(table, i, 3, range_str)
                self._set_cell(table, i, 4, cond.status.upper())
            
            self.doc.add_paragraph()
        
        # Test location and time
        table = self.doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'
        
        loc_data = [
            ('Test Location', conditions.test_location),
            ('Location Detail', conditions.location_detail),
            ('Test Date', conditions.test_date.strftime("%d %B %Y") if conditions.test_date else 'N/A'),
            ('Start Time', conditions.start_time or 'N/A'),
            ('End Time', conditions.end_time or 'N/A'),
        ]
        
        for i, (label, value) in enumerate(loc_data):
            self._set_cell(table, i, 0, label, bold=True)
            self._set_cell(table, i, 1, value)
        
        self.doc.add_paragraph()
    
    def _build_equipment(self, equipment: TestEquipment):
        """Build test equipment section."""
        self.doc.add_heading('4. TEST EQUIPMENT', level=1)
        
        if equipment.items:
            table = self.doc.add_table(rows=len(equipment.items) + 1, cols=6)
            table.style = 'Table Grid'
            
            # Header
            headers = ['Equipment ID', 'Name', 'Type', 'Serial Number', 'Cal. Date', 'Cal. Valid Until']
            for i, header in enumerate(headers):
                self._set_cell(table, 0, i, header, bold=True, bg_color=DocxColors.TABLE_HEADER, font_color=DocxColors.WHITE)
            
            # Data
            for i, item in enumerate(equipment.items, 1):
                self._set_cell(table, i, 0, item.equipment_id)
                self._set_cell(table, i, 1, item.name)
                self._set_cell(table, i, 2, item.equipment_type)
                self._set_cell(table, i, 3, item.serial_number)
                self._set_cell(table, i, 4, item.calibration_date.strftime("%d %b %Y") if item.calibration_date else 'N/A')
                self._set_cell(table, i, 5, item.calibration_valid_until.strftime("%d %b %Y") if item.calibration_valid_until else 'N/A')
        else:
            self.doc.add_paragraph("No equipment records available.")
        
        self.doc.add_paragraph()
    
    def _build_test_results(self, results: list[TestResult]):
        """Build test results section."""
        self.doc.add_heading('5. TEST PROCEDURES AND RESULTS', level=1)
        
        for i, result in enumerate(results, 1):
            self._build_single_test_result(result, i)
    
    def _build_single_test_result(self, result: TestResult, index: int):
        """Build a single test result subsection."""
        self.doc.add_heading('5.{} {} {}'.format(index, result.procedure.test_code, result.procedure.test_name), level=2)
        
        if result.procedure.purpose:
            p = self.doc.add_paragraph()
            p.add_run('Purpose: ').bold = True
            p.add_run(result.procedure.purpose)
        
        # Observations
        if result.observations:
            self.doc.add_heading('5.{}.1 Observations'.format(index), level=3)
            
            table = self.doc.add_table(rows=len(result.observations) + 1, cols=4)
            table.style = 'Table Grid'
            
            # Header
            headers = ['Obs. #', 'Value', 'Unit', 'Notes']
            for i, header in enumerate(headers):
                self._set_cell(table, 0, i, header, bold=True, bg_color=DocxColors.TABLE_HEADER, font_color=DocxColors.WHITE)
            
            # Data
            for i, obs in enumerate(result.observations, 1):
                self._set_cell(table, i, 0, str(obs.observation_number))
                self._set_cell(table, i, 1, str(obs.value))
                self._set_cell(table, i, 2, obs.unit)
                self._set_cell(table, i, 3, obs.notes or '')
            
            self.doc.add_paragraph()
        
        # Results summary (OBSERVED / CALCULATED / LIMIT / RESULT)
        self.doc.add_heading('5.{}.2 Results Summary'.format(index), level=3)
        
        # Build summary table
        rows = []
        
        if result.mean is not None:
            rows.append(['Mean', '{:.6f}'.format(result.mean), '', '', ''])
        
        if result.standard_deviation is not None:
            rows.append([
                'Standard Deviation',
                '{:.6f}'.format(result.standard_deviation),
                '{:.4f} d'.format(result.calculated_in_d) if result.calculated_in_d else '',
                '{} {}'.format(result.limit_value, result.limit_unit) if result.limit_value else '',
                result.status.upper() if result.status else '',
            ])
        
        if result.absolute_error is not None:
            rows.append(['Absolute Error', '{:.6f}'.format(result.absolute_error), '', '', ''])
        
        if result.deviation_from_reference is not None:
            rows.append(['Deviation from Reference', '{:.6f}'.format(result.deviation_from_reference), '', '', ''])
        
        if rows:
            table = self.doc.add_table(rows=len(rows) + 1, cols=5)
            table.style = 'Table Grid'
            
            # Header
            headers = ['Parameter', 'Observed', 'Calculated', 'Limit', 'Result']
            for i, header in enumerate(headers):
                self._set_cell(table, 0, i, header, bold=True, bg_color=DocxColors.TABLE_HEADER, font_color=DocxColors.WHITE)
            
            # Data with color coding
            for i, row in enumerate(rows, 1):
                for j, cell_value in enumerate(row):
                    bg_color = None
                    if j == 1:  # Observed
                        bg_color = DocxColors.OBSERVED_BG
                    elif j == 2:  # Calculated
                        bg_color = DocxColors.CALCULATED_BG
                    elif j == 3:  # Limit
                        bg_color = DocxColors.LIMIT_BG
                    elif j == 4:  # Result
                        bg_color = DocxColors.RESULT_BG
                    
                    self._set_cell(table, i, j, cell_value, bg_color=bg_color)
        
        # Status
        if result.status:
            p = self.doc.add_paragraph()
            p.add_run('Result: ').bold = True
            run = p.add_run(result.status.upper())
            if result.status == 'pass':
                run.font.color.rgb = RGBColor(0x22, 0x8b, 0x22)
            elif result.status == 'fail':
                run.font.color.rgb = RGBColor(0xdc, 0x14, 0x3c)
        
        if result.reason:
            p = self.doc.add_paragraph()
            run = p.add_run(result.reason)
            run.italic = True
            run.font.size = Pt(9)
        
        self.doc.add_paragraph()
    
    def _build_compliance(self, compliance: ComplianceResult):
        """Build compliance section."""
        self.doc.add_heading('6. COMPLIANCE RESULT', level=1)
        
        # Overall status
        table = self.doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        self._set_cell(table, 0, 0, 'Overall Compliance Status', bold=True)
        
        bg_color = DocxColors.PASS_BG if compliance.overall_status == 'pass' else DocxColors.FAIL_BG
        self._set_cell(table, 0, 1, compliance.overall_status.upper(), bold=True, bg_color=bg_color)
        
        self.doc.add_paragraph()
        
        # Individual test results summary
        if compliance.test_results:
            table = self.doc.add_table(rows=len(compliance.test_results) + 1, cols=5)
            table.style = 'Table Grid'
            
            # Header
            headers = ['Test', 'Parameter', 'Calculated', 'Limit', 'Status']
            for i, header in enumerate(headers):
                self._set_cell(table, 0, i, header, bold=True, bg_color=DocxColors.TABLE_HEADER, font_color=DocxColors.WHITE)
            
            # Data
            for i, tr in enumerate(compliance.test_results, 1):
                self._set_cell(table, i, 0, tr.procedure.test_name)
                self._set_cell(table, i, 1, 'Standard Deviation' if tr.standard_deviation else '')
                self._set_cell(table, i, 2, '{:.4f} d'.format(tr.calculated_in_d) if tr.calculated_in_d else '')
                self._set_cell(table, i, 3, '{} {}'.format(tr.limit_value, tr.limit_unit) if tr.limit_value else '')
                self._set_cell(table, i, 4, tr.status.upper() if tr.status else '')
        
        if compliance.remarks:
            p = self.doc.add_paragraph()
            p.add_run('Remarks: ').bold = True
            p.add_run(compliance.remarks)
        
        self.doc.add_paragraph()
    
    def _build_remarks(self, remarks: str):
        """Build remarks section."""
        self.doc.add_heading('7. REMARKS', level=1)
        self.doc.add_paragraph(remarks)
        self.doc.add_paragraph()
    
    def _build_signatures(self, report: TestReport):
        """Build signatures section."""
        self.doc.add_heading('8. SIGNATURES', level=1)
        
        sig_data = [['Role', 'Name', 'Title', 'Date']]
        
        if report.technician:
            sig_data.append([
                'Technician',
                report.technician.name,
                report.technician.title,
                report.technician.date.strftime("%d %B %Y") if report.technician.date else '',
            ])
        
        if report.reviewer:
            sig_data.append([
                'Reviewer',
                report.reviewer.name,
                report.reviewer.title,
                report.reviewer.date.strftime("%d %B %Y") if report.reviewer.date else '',
            ])
        
        if report.approver:
            sig_data.append([
                'Approver',
                report.approver.name,
                report.approver.title,
                report.approver.date.strftime("%d %B %Y") if report.approver.date else '',
            ])
        
        if len(sig_data) > 1:
            table = self.doc.add_table(rows=len(sig_data), cols=4)
            table.style = 'Table Grid'
            
            for i, row in enumerate(sig_data):
                for j, cell_value in enumerate(row):
                    bold = (i == 0)
                    self._set_cell(table, i, j, cell_value, bold=bold)
        
        self.doc.add_paragraph()
    
    def _build_disclaimer(self):
        """Build disclaimer section."""
        disclaimer_text = (
            "DISCLAIMER: This report is generated by software for informational purposes only. "
            "It does not constitute legal certification, regulatory approval, or official accreditation. "
            "Actual compliance determination must be made by authorized personnel in accordance with applicable regulations. "
            "The values presented are based on measurements and calculations performed using the described methodology."
        )
        
        p = self.doc.add_paragraph()
        run = p.add_run(disclaimer_text)
        run.italic = True
        run.font.size = Pt(8)
        run.font.color.rgb = DocxColors.TEXT_SECONDARY
    
    def _add_header_footer(self, report: TestReport):
        """Add header and footer to the document."""
        section = self.doc.sections[0]
        
        # Header
        header = section.header
        header_para = header.paragraphs[0]
        header_para.text = "Report: {} | OIML R-76 ({})".format(
            report.identification.report_number,
            report.identification.standard_version
        )
        header_para.style.font.size = Pt(8)
        
        # Footer
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = "Generated: {} | CONFIDENTIAL".format(
            datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S UTC")
        )
        footer_para.style.font.size = Pt(7)
    
    def _set_cell(self, table, row, col, text, bold=False, bg_color=None, font_color=None):
        """Set cell text with formatting."""
        cell = table.cell(row, col)
        cell.text = ''
        
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.bold = bold
        
        if font_color:
            run.font.color.rgb = font_color
        
        if bg_color:
            shading_elm = parse_xml(
                '<w:shd {} w:fill="{}"/>'.format(
                    nsdecls('w'),
                    '{:02x}{:02x}{:02x}'.format(bg_color[0], bg_color[1], bg_color[2])
                )
            )
            cell._tc.get_or_add_tcPr().append(shading_elm)
